# -*- coding: utf-8 -*-
"""
Different Apache license headers for different types of files
# Copyright 2015 Glenn ten Cate, Riccardo ten Cate
Licensed under the Apache License, Version 2.0 (the "License");
you may not use this file except in compliance with the License.
You may obtain a copy of the License at
http://www.apache.org/licenses/LICENSE-2.0
Unless required by applicable law or agreed to in writing, software
distributed under the License is distributed on an "AS IS" BASIS,
WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
See the License for the specific language governing permissions and
limitations under the License.
"""

import os, markdown, datetime, string, base64
from OpenSSL import SSL, rand
from docx import Document
from BeautifulSoup import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from functools import wraps 
from sqlite3 import dbapi2 as sqlite3
from flask import Flask, request, session, g, redirect, url_for, abort, \
     render_template, flash, Markup, make_response

# create the application
app = Flask(__name__)

def add_response_headers(headers={}):
    """This decorator adds the headers passed in to the response"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            resp = make_response(f(*args, **kwargs))
            h = resp.headers
            for header, value in headers.items():
                h[header] = value
            return resp
        return decorated_function
    return decorator

def security(f):
    """This decorator passes multiple security headers"""
    return add_response_headers({'X-Frame-Options': 'deny', 'X-XSS-Protection': '1', 'X-Content-Type-Options': 'nosniff', 'Cache-Control': 'no-store, no-cache','Strict-Transport-Security': 'max-age=16070400; includeSubDomains', 'Server': 'Security Knowledge Framework'})(f)

def check_token():
    """Checks the submitted CSRF token"""
    if not session.get('csrf_token') == request.form['csrf_token']:
        session.destroy()
        return abort(500)(f)

def generate_pass():
    chars = string.letters + string.digits + '+/'
    assert 256 % len(chars) == 0  # non-biased later modulo
    PWD_LEN = 12
    password = ''.join(chars[ord(c) % len(chars)] for c in os.urandom(PWD_LEN))
    return password

#secret key for flask internal session use
secret_key = rand.bytes(512)
#password = generate_pass()

# Load default config and override config from an environment variable
# You can also replace password with static password:  PASSWORD='pass!@#example'
app.config.update(dict(
    DATABASE=os.path.join(app.root_path, 'skf.db'),
    DEBUG=True,
    SECRET_KEY=secret_key,
    USERNAME='admin',
    SESSION_COOKIE_SECURE=False,
    PASSWORD='default'
))


def connect_db():
    """Connects to the specific database."""
    rv = sqlite3.connect(app.config['DATABASE'])
    rv.row_factory = sqlite3.Row
    return rv


def init_db():
    """Initializes the database."""
    db = get_db()
    with app.open_resource('schema.sql', mode='r') as f:
        db.cursor().executescript(f.read())
    db.commit()


@app.cli.command('initdb')
def initdb_command():
    """Creates the database tables."""
    init_db()
    print('Initialized the database.')


def get_db():
    """Opens a new database connection if there is none yet for the
    current application context.
    """
    if not hasattr(g, 'sqlite_db'):
        g.sqlite_db = connect_db()
    return g.sqlite_db

def get_filepaths(directory):
    """
    This function will generate the file names in a directory 
    tree by walking the tree either top-down or bottom-up. For each 
    directory in the tree rooted at directory top (including top itself), 
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = [] 
    for root, directories, files in os.walk(directory):
        for filename in files:
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)
    return file_paths  

def get_num(x):
    """get numbers from a string"""
    return int(''.join(ele for ele in x if ele.isdigit()))

@app.teardown_appcontext
def close_db(error):
    """Closes the database again at the end of the request."""
    if hasattr(g, 'sqlite_db'):
        g.sqlite_db.close()

def projects_functions_techlist():
    """get list of technology used for creating project functions"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    cur = db.execute('SELECT techID, techName, vulnID from techhacks ORDER BY techID DESC')
    entries = cur.fetchall()
    return entries 

@app.route('/')
@security
def show_landing():
    """show the loging page and set default code language"""
    return render_template('login.html')

@app.route('/dashboard', methods=['GET'])
@security
def dashboard():
    """show the landing page"""
    if not session.get('logged_in'):
        abort(401)
    return render_template('dashboard.html')

@app.route('/login', methods=['GET', 'POST'])
@security
def login():
    """validate the login data for access dashboard page"""
    error = None
    csrf_token_raw = rand.bytes(128)
    csrf_token = base64.b64encode(csrf_token_raw)
    if request.method == 'POST':
        if request.form['username'] != app.config['USERNAME']:
            error = 'Invalid username/password'
        elif request.form['password'] != app.config['PASSWORD']:
            error = 'Invalid username/password'
        else:
            session['logged_in'] = True
            session['csrf_token'] = csrf_token
            session['code_lang'] = "php"
            return render_template('dashboard.html')
    return render_template('login.html', error=error)

@app.route('/logout', methods=['GET', 'POST'])
@security
def logout():
    """logout and destroy session"""
    session.pop('logged_in', None)
    return redirect(url_for('login'))

@app.route('/code/<code_lang>', methods=['GET'])
@security
def set_code_lang(code_lang):
    """set a code language: php java python perl"""
    if not session.get('logged_in'):
        abort(401)
    allowed = "php java python perl"
    found = allowed.find(code_lang)
    if found != -1:
        session['code_lang'] = code_lang
    return redirect(url_for('code_examples'))

@app.route('/code-examples', methods=['GET'])
@security
def code_examples():
    """Shows the knowledge base markdown files."""
    if not session.get('logged_in'):
        abort(401)
    items = []
    id_items = []
    full_file_paths = []
    allowed = set(string.ascii_lowercase + string.ascii_uppercase + '.')
    if set(session['code_lang']) <= allowed:
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/code_examples/"+session['code_lang']))
        for path in full_file_paths:
            id_item = get_num(path)
            path = path.split("-")
            y = len(path)-3 
            kb_name_uri = path[(y)]
            kb_name = kb_name_uri.replace("_", " ")
            items.append(kb_name)
            id_items.append(id_item)
    return render_template('code-examples.html', items=items, id_items=id_items)

@app.route('/code-search', methods=['POST'])
@security
def show_code_search():
    """show the landing page"""
    if not session.get('logged_in'):
        abort(401)
    search = request.form['search']
    full_file_paths = []
    allowed = set(string.ascii_lowercase + string.ascii_uppercase + '.')
    if set(session['code_lang']) <= allowed:
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/code_examples/"+session['code_lang']))
        for path in full_file_paths:
            found = path.find(search)
            if found != -1:
                filemd = open(path, 'r').read()
                content = Markup(markdown.markdown(filemd))
                path = path.split("-")
                y = len(path)-3
                kb_name_uri = path[(y)]
                kb_name = kb_name_uri.replace("_", " ")
    return render_template('code-examples-search.html', **locals())

@app.route('/code-item', methods=['POST'])
@security
def show_code_item():
    """show the coding examples page"""
    if not session.get('logged_in'):
        abort(401)
    id = int(request.form['id'])
    items = []
    full_file_paths = []
    allowed = set(string.ascii_lowercase + string.ascii_uppercase + '.')
    if set(session['code_lang']) <= allowed:
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/code_examples/"+session['code_lang']))
        for path in full_file_paths:
            if id == get_num(path):
                filemd = open(path, 'r').read()
                content = Markup(markdown.markdown(filemd)) 
    return render_template('code-examples-item.html', **locals())

@app.route('/kb-search', methods=['POST'])
@security
def show_kb_search():
    """show the knowledge base search page"""
    if not session.get('logged_in'):
        abort(401)
    search = request.form['search']
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
    for path in full_file_paths:
        found = path.find(search)
        if found != -1:
            filemd = open(path, 'r').read()
            content = Markup(markdown.markdown(filemd))
            path = path.split("-")
            y = len(path)-3
            kb_name_uri = path[(y)]
            kb_name = kb_name_uri.replace("_", " ")
    return render_template('knowledge-base-search.html', **locals())


@app.route('/kb-item', methods=['POST'])
@security
def show_kb_item():
    """show the knowledge base search result page"""
    if not session.get('logged_in'):
        abort(401)
    id = int(request.form['id'])
    items = []
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown"))
    for path in full_file_paths:
        if id == get_num(path):
            filemd = open(path, 'r').read()
            content = Markup(markdown.markdown(filemd)) 
    return render_template('knowledge-base-item.html', **locals())

@app.route('/knowledge-base', methods=['GET'])
@security
def knowledge_base():
    """Shows the knowledge base markdown files."""
    if not session.get('logged_in'):
        abort(401)
    items = []
    id_items = []
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
    for path in full_file_paths:
        id_item = get_num(path)
        path = path.split("-")
        y = len(path)-3 
        kb_name_uri = path[(y)]
        kb_name = kb_name_uri.replace("_", " ")
        items.append(kb_name)
        id_items.append(id_item)
    return render_template('knowledge-base.html', items=items, id_items=id_items)

@app.route('/project-new', methods=['GET'])
@security
def projects():
    """show the create new project page"""
    if not session.get('logged_in'):
        abort(401)
    return render_template('project-new.html', csrf_token=session['csrf_token'])

@app.route('/project-add', methods=['POST'])
@security
def add_entry():
    """add a new project to database"""
    if not session.get('logged_in'):
        abort(401)
    check_token()
    db = get_db()
    date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    db.execute('INSERT INTO projects (timestamp, projectName, projectVersion, projectDesc) VALUES (?, ?, ?, ?)',
               [date, request.form['inputName'], request.form['inputVersion'], request.form['inputDesc']])
    db.commit()
    return redirect(url_for('project_list'))

@app.route('/project-del', methods=['POST'])
@security
def project_del():
    """delete project from database"""
    if not session.get('logged_in'):
        abort(401)
    check_token()
    db = get_db()
    db.execute("DELETE FROM projects WHERE projectID=?",
               [request.form['projectID']])
    db.commit()
    return render_template('reload.html')

@app.route('/project-list', methods=['GET'])
@security
def project_list():
    """show the project list page"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    cur = db.execute('SELECT projectName, projectVersion, projectDESC, projectID, timestamp FROM projects ORDER BY projectID DESC')
    entries = cur.fetchall()
    return render_template('project-list.html', entries=entries, csrf_token=session['csrf_token'])

@app.route('/project-options/<project_id>', methods=['GET'])
@security
def projects_options(project_id):
    """show the project options landing page"""
    if not session.get('logged_in'):
        abort(401)
    return render_template('project-options.html', project_id=project_id, csrf_token=session['csrf_token'])

@app.route('/project-functions/<project_id>', methods=['GET'])
@security
def project_functions(project_id):
    """show the lproject functions page"""
    if not session.get('logged_in'):
        abort(401)
    techlist = projects_functions_techlist()
    db = get_db()
    db.commit()
    cur = db.execute('SELECT paramID, functionName, functionDesc, projectID, tech, entryDate FROM parameters WHERE projectID=? ORDER BY projectID DESC',
                      [project_id])
    entries = cur.fetchall()
    return render_template('project-functions.html', project_id=project_id, techlist=projects_functions_techlist(), entries=entries, csrf_token=session['csrf_token'])

@app.route('/project-function-del', methods=['POST'])
@security
def function_del():
    """delete a project function"""
    if not session.get('logged_in'):
        abort(401)
    check_token()
    id = int(request.form['projectID'])
    db = get_db()
    db.execute("DELETE FROM parameters WHERE projectID=? AND paramID=?",
               [request.form['projectID'],request.form['paramID']])
    db.commit()
    redirect_url = "/project-functions/"+str(id)
    return redirect(redirect_url)


@app.route('/project-function-add', methods=['POST'])
@security
def add_function():
    """add a project function"""
    if not session.get('logged_in'):
        abort(401)
    check_token()
    id = int(request.form['project_id'])
    f = request.form
    for key in f.keys():
        for value in f.getlist(key):
                found = key.find("test")
                if found != -1:
                    db = get_db()
                    date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    db.execute('INSERT INTO parameters (entryDate, functionName, functionDesc, tech, projectID) VALUES (?, ?, ?, ?, ?)',
                           [date, request.form['functionName'], request.form['functionDesc'], value, request.form['project_id']])
                    db.commit()
    redirect_url = '/project-functions/'+str(id)
    return redirect(redirect_url)

@app.route('/project-checklists/<project_id>', methods=['GET'])
@security
def project_checklists(project_id):
    """show the project checklists page"""
    if not session.get('logged_in'):
        abort(401)
    csrf_token=session['csrf_token']
    db = get_db()
    cur = db.execute('SELECT projectName FROM projects WHERE projectID=?',
                        [project_id])
    row = cur.fetchall()
    prep = row[0]
    projectName = prep[0]
    owasp_items = []
    owasp_ids = []
    owasp_kb_ids = []
    owasp_content = []
    owasp_items_lvl1 = []
    owasp_ids_lvl1 = []
    owasp_kb_ids_lvl1 = []
    owasp_content_lvl1 = []
    owasp_items_lvl2 = []
    owasp_ids_lvl2 = []
    owasp_kb_ids_lvl2 = []
    owasp_content_lvl2 = []
    owasp_items_lvl3 = []
    owasp_ids_lvl3 = []
    owasp_kb_ids_lvl3 = []
    owasp_content_lvl3 = []
    custom_items = []
    custom_ids = []
    custom_kb_ids = []
    custom_content = []
    basic_items = []
    basic_ids = []
    basic_kb_ids = []
    basic_content = []
    advanced_items = []
    advanced_ids = []
    advanced_kb_ids = []
    advanced_content = []
    full_file_paths = []
    full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/checklists"))
    for path in full_file_paths:
       found = path.find("owasp")
       if found != -1:
            owasp_org_path = path
            owasp_list = "owasp"
            owasp_path = path.split("-")
            owasp_kb = owasp_path[5]
            owasp_checklist_name = owasp_path[3]
            owasp_id = get_num(owasp_path[1])
            owasp_items.append(owasp_checklist_name)
            owasp_ids.append(owasp_id)
            owasp_kb_ids.append(owasp_kb)
            filemd = open(owasp_org_path, 'r').read()
            owasp_content.append(Markup(markdown.markdown(filemd)))
    for path in full_file_paths:
       found = path.find("ASVS-level-1")
       if found != -1:
            owasp_org_path = path
            owasp_list = "ASVS-level-1"
            owasp_path_lvl1 = path.split("-")
            owasp_kb = owasp_path_lvl1[7]
            owasp_checklist_name = owasp_path_lvl1[3] +" "+owasp_path_lvl1[4]+" "+owasp_path_lvl1[5]
            owasp_id = get_num(owasp_path[1])
            owasp_items_lvl1.append(owasp_checklist_name)
            owasp_ids_lvl1.append(owasp_id)
            owasp_kb_ids_lvl1.append(owasp_kb)
            filemd = open(owasp_org_path, 'r').read()
            owasp_content_lvl1.append(Markup(markdown.markdown(filemd)))
    for path in full_file_paths:
       found = path.find("ASVS-level-2")
       if found != -1:
            owasp_org_path = path
            owasp_list = "ASVS-level-2"
            owasp_path_lvl2 = path.split("-")
            owasp_kb = owasp_path_lvl2[7]
            owasp_checklist_name = owasp_path_lvl2[3] +" "+owasp_path_lvl2[4]+" "+owasp_path_lvl2[5]
            owasp_id = get_num(owasp_path[1])
            owasp_items_lvl2.append(owasp_checklist_name)
            owasp_ids_lvl2.append(owasp_id)
            owasp_kb_ids_lvl2.append(owasp_kb)
            filemd = open(owasp_org_path, 'r').read()
            owasp_content_lvl2.append(Markup(markdown.markdown(filemd)))
    for path in full_file_paths:
       found = path.find("ASVS-level-3")
       if found != -1:
            owasp_org_path = path
            owasp_list = "ASVS-level-3"
            owasp_path_lvl3 = path.split("-")
            owasp_kb = owasp_path_lvl3[7]
            owasp_checklist_name = owasp_path_lvl3[3] +" "+owasp_path_lvl3[4]+" "+owasp_path_lvl3[5]
            owasp_id = get_num(owasp_path[1])
            owasp_items_lvl3.append(owasp_checklist_name)
            owasp_ids_lvl3.append(owasp_id)
            owasp_kb_ids_lvl3.append(owasp_kb)
            filemd = open(owasp_org_path, 'r').read()
            owasp_content_lvl3.append(Markup(markdown.markdown(filemd)))
    for path in full_file_paths:
       found = path.find("CS_basic_audit")
       if found != -1:
            basic_org_path = path
            basic_list = "CS_basic_audit"
            basic_path = path.split("-")
            basic_kb = basic_path[5]
            basic_checklist_name = basic_path[3]
            basic_id = get_num(basic_path[1])
            basic_items.append(basic_checklist_name)
            basic_ids.append(basic_id)
            basic_kb_ids.append(basic_kb)
            filemd = open(basic_org_path, 'r').read()
            basic_content.append(Markup(markdown.markdown(filemd)))
    for path in full_file_paths:
       found = path.find("CS_advanced_audit")
       if found != -1:
            advanced_org_path = path
            advanced_list = "CS_advanced_audit"
            advanced_path = path.split("-")
            advanced_kb = advanced_path[5]
            advanced_name = advanced_path[3]
            advanced_id = get_num(advanced_path[1])
            advanced_items.append(advanced_name)
            advanced_ids.append(advanced_id)
            advanced_kb_ids.append(advanced_kb)
            filemd = open(advanced_org_path, 'r').read()
            advanced_content.append(Markup(markdown.markdown(filemd)))
    for path in full_file_paths:
       found = path.find("custom")
       if found != -1:
            custom_org_path = path
            custom_list = "custom"
            custom_path = path.split("-")
            custom_kb = custom_path[5]
            custom_name = custom_path[3]
            custom_id = get_num(custom_path[1])
            custom_items.append(custom_name)
            custom_ids.append(custom_id)
            custom_kb_ids.append(custom_kb)
            filemd = open(custom_org_path, 'r').read()
            custom_content.append(Markup(markdown.markdown(filemd)))
    return render_template('project-checklists.html', **locals())

@app.route('/project-checklist-add', methods=['POST'])
@security
def add_checklist():
    """add project checklist"""
    if not session.get('logged_in'):
        abort(401)
    check_token()
    f = request.form
    i = 0
    for key in f.keys():
        for value in f.getlist(key):
            found = key.find("vuln")
            if found != -1:
                listID = "listID"+str(i)
                answerID = "answer"+str(i)
                questionID = "questionID"+str(i) 
                vulnID = "vulnID"+str(i)
                date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                db = get_db()
                db.execute('INSERT INTO questionlist (entryDate, answer, projectName, projectID, questionID, vulnID, listName) VALUES (?, ?, ?, ?, ?, ?, ?)',
                           [date, request.form[answerID], request.form['projectName'], request.form['projectID'], request.form[questionID], request.form[vulnID], request.form[listID]])
                db.commit()
                i += 1
    redirect_url = "/results-checklists"
    return redirect(redirect_url)

@app.route('/results-checklists', methods=['GET'])
@security
def results_checklists():
    """show the results checklists page"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    cur = db.execute('SELECT q.answer, q.projectID, q.questionID,  q.vulnID, q.listName, q.entryDate, p.projectName, p.projectVersion, p.projectDesc FROM questionlist AS q JOIN projects AS p ON q.projectID = p.projectID  GROUP BY q.listName, q.entryDate ORDER BY p.projectName ASC')
    entries = cur.fetchall()
    return render_template('results-checklists.html', entries=entries, csrf_token=session['csrf_token'])

@app.route('/results-functions', methods=['GET'])
@security
def results_functions():
    """show the results functions page"""
    if not session.get('logged_in'):
        abort(401)
    db = get_db()
    cur = db.execute('SELECT p.projectName, p.projectID, par.entryDate, p.projectDesc, p.projectVersion, par.paramID, par.functionName, par.projectID FROM projects AS p join parameters AS par on p.projectID = par.projectID GROUP BY p.projectVersion ')
    entries = cur.fetchall()
    return render_template('results-functions.html', entries=entries, csrf_token=session['csrf_token'])

@app.route('/results-functions-del', methods=['POST'])
@security
def functions_del():
    """delete functions result items"""
    if not session.get('logged_in'):
        abort(401)
    check_token()
    db = get_db()
    db.execute("DELETE FROM parameters WHERE entryDate=?",
               [request.form['entryDate']])
    db.commit()
    return render_template('reload.html')

@app.route('/results-checklists-del', methods=['POST'])
@security
def checklists_del():
    """delete checklist result item"""
    if not session.get('logged_in'):
        abort(401)
    check_token()
    db = get_db()
    db.execute("DELETE FROM questionlist WHERE entryDate=?",
               [request.form['entryDate']])
    db.commit()
    return render_template('reload.html')


@app.route('/results-checklist-report/<entryDate>', methods=['GET'])
@security
def checklist_results(entryDate):
    """show checklist results report"""
    if not session.get('logged_in'):
        abort(401)
    id_items = []
    content = []
    full_file_paths = []
    db = get_db()
    cur = db.execute("SELECT * FROM questionlist WHERE answer='no' AND entryDate=?",
               [entryDate])
    entries = cur.fetchall()
    for entry in entries:
        projectName = entry[3]
        vulnID = entry[5]
        listName = entry[6]
        entryDate = entry[7]
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for path in full_file_paths:
            org_path = path
            path_vuln = get_num(path)
            if int(vulnID) == int(path_vuln):
                print path_vuln

                filemd = open(org_path, 'r').read()
                content.append(Markup(markdown.markdown(filemd)))
    return render_template('results-checklist-report.html', **locals())


@app.route('/results-checklist-docx/<entryDate>')
def download_file_checklist(entryDate):
    """Download checklist results report in docx"""
    if not session.get('logged_in'):
        abort(401)
    content_raw = []
    content_title = []
    db = get_db()
    cur = db.execute("SELECT * FROM questionlist WHERE answer='no' AND entryDate=?",
               [entryDate])
    entries = cur.fetchall()
    document = Document()
    document.add_picture('static/img/logo.png', width=Inches(1.25))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('Security Knowledge Framework', 0)
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    projectName = entries[0][3]
    listName = entries[0][6]
    p.add_run('Used Checklist: '+listName)
    p.add_run('\r\n')
    p.add_run('Date: '+datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    p.add_run('\r\n')
    p.add_run('Project: '+projectName)
    document.add_page_break()
    p = document.add_heading('Table of contents', level=1)
    p.add_run('\r\n')
    document.add_paragraph('Introduction')
    for entry in entries:
        projectName = entry[3]
        vulnID = entry[5]
        listName = entry[6]
        entryDate = entry[7]
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for path in full_file_paths:
            org_path = path
            path = path.split("-")
            name = org_path[3]
            path_vuln = get_num(path[1])
            if int(vulnID) == int(path_vuln):
                print "\n"
                print org_path
                filemd = open(org_path, 'r').read()
                content = Markup(markdown.markdown(filemd))
                text = ''.join(BeautifulSoup(content).findAll(text=True))
                text_encode = text.encode('utf-8')
                content_title.append(text_encode.splitlines()[0])
                text_encode = text_encode.replace("Solution", "\nSolution");
                content_raw.append(text_encode)
    for item in content_title:
        p = document.add_paragraph(item)
        p.add_run()
    document.add_page_break()
    document.add_heading('Introduction', level=1)
    p = document.add_paragraph(
        'The security knowledge framework is composed by means of the highest security standards currently available and is designed to maintain the integrety of your application, so you and your costumers sensitive data is protected against hackers. This document is provided with a checklist in which the programmers of your application had to run through in order to provide a secure product.'
    )
    p.add_run('\n')
    p = document.add_paragraph(
        'In the post-development stage of the security knowledge framework the developer double-checks his application against a checklist which consists out of several questions asking the developer about different stages of development and the methodology of implementing different types of functionality the application contains. After filling in this checklist the developer gains feedback on the failed checklist items providing him with solutions about how to solve the additional vulnerability\'s found in the application.'
    )
    document.add_page_break()
    i = 0
    for item in content_raw:
        document.add_heading(content_title[i], level=1)
        p = document.add_paragraph(item.partition("\n")[2])
        p.add_run("\n")
        document.add_page_break()
        i += 1
    document.save('checklist-security-report.docx')
    headers = {"Content-Disposition": "attachment; filename=%s" % "checklist-security-report.docx"}
    with open("checklist-security-report.docx", 'r') as f:
        body = f.read()
    return make_response((body, headers))
    
    
@app.route('/results-function-report/<entryDate>', methods=['GET'])
@security
def function_results(entryDate):
    """show checklist results report"""
    if not session.get('logged_in'):
        abort(401)
    id_items = []
    content = []
    full_file_paths = []
    db = get_db()
    cur = db.execute("SELECT projects.projectName, projects.projectID, projects.projectVersion, parameters.functionName, parameters.tech, parameters.functionDesc, parameters.entryDate FROM projects JOIN parameters ON parameters.projectID=projects.projectID WHERE parameters.entryDate=?",
               [entryDate])
    entries = cur.fetchall()
    for entry in entries:
        projectName = entry[0]
        vulnID = entry[4]
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for path in full_file_paths:
            org_path = path
            path_vuln = get_num(path)
            if int(vulnID) == int(path_vuln):
                filemd = open(org_path, 'r').read()
                content.append(Markup(markdown.markdown(filemd)))
    return render_template('results-function-report.html', **locals())

@app.route('/results-function-docx/<entryDate>')
def download_file_function(entryDate):
    """Download checklist results report in docx"""
    if not session.get('logged_in'):
        abort(401)
    content_raw = []
    content_title = []
    db = get_db()
    cur = db.execute("SELECT projects.projectName, projects.projectID, projects.projectVersion, parameters.functionName, parameters.tech, parameters.functionDesc, parameters.entryDate FROM projects JOIN parameters ON parameters.projectID=projects.projectID WHERE parameters.entryDate=?",
               [entryDate])
    entries = cur.fetchall()
    document = Document()
    document.add_picture('static/img/logo.png', width=Inches(2.25), height=Inches(2.25))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('Security Knowledge Framework', 0)
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    projectName = entries[0][0]
    functionName = entries[0][3]
    functionDesc= entries[0][5]
    p.add_run('Function Name: '+functionName)
    p.add_run('Function Description: '+functionDesc)
    p.add_run('\r\n')
    p.add_run('Date: '+datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    p.add_run('\r\n')
    p.add_run('Project: '+projectName)
    document.add_page_break()
    p = document.add_heading('Table of contents', level=1)
    p.add_run('\r\n')
    document.add_paragraph('Introduction')
    for entry in entries:
        print entry
        entryDate = entry[6]
        vulnID = entry[4]
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for path in full_file_paths:
            org_path = path
            path_vuln = get_num(path)
            if int(vulnID) == int(path_vuln):
                filemd = open(org_path, 'r').read()
                content = Markup(markdown.markdown(filemd))
                text = ''.join(BeautifulSoup(content).findAll(text=True))
                text_encode = text.encode('utf-8')
                content_title.append(text_encode.splitlines()[0])
                text_encode = text_encode.replace("Solution", "\nSolution");
                content_raw.append(text_encode)
    for item in content_title:
        p = document.add_paragraph(item)
        p.add_run()
    document.add_page_break()
    document.add_heading('Introduction', level=1)
    p = document.add_paragraph(
        'The security knowledge framework is composed by means of the highest security standards currently available and is designed to maintain the integrety of your application, so you and your costumers sensitive data is protected against hackers. This document is provided with a checklist in which the programmers of your application had to run through in order to provide a secure product.'
    )
    p.add_run('\n')
    p = document.add_paragraph(
        'In the post-development stage of the security knowledge framework the developer double-checks his application against a checklist which consists out of several questions asking the developer about different stages of development and the methodology of implementing different types of functionality the application contains. After filling in this checklist the developer gains feedback on the failed checklist items providing him with solutions about how to solve the additional vulnerability\'s found in the application.'
    )
    document.add_page_break()
    i = 0
    for item in content_raw:
        document.add_heading(content_title[i], level=1)
        p = document.add_paragraph(item.partition("\n")[2])
        p.add_run("\n")
        document.add_page_break()
        i += 1
    document.save('function-security-report.docx')
    headers = {"Content-Disposition": "attachment; filename=%s" % "function-security-report.docx"}
    with open("function-security-report.docx", 'r') as f:
        body = f.read()
    return make_response((body, headers))

if __name__ == "__main__":
     if os.path.isfile('server.crt') == False: 
        #print "Generated Password for access SKF: "+password
        app.run(host='127.0.0.1', port=5443, ssl_context='adhoc')
     else:
        context = SSL.Context(SSL.TLSv1_METHOD)
        context.use_privatekey_file('server.key')  #Location of Key
        context.use_certificate_file('server.crt') #Location of Cert
        context.set_cipher_list('TLSv1+HIGH:!aNULL:!eNULL:!3DES:@STRENGTH')
        #print "Generated Password for access SKF: "+password
        app.run(host='127.0.0.1', port=5443, ssl_context=context)
